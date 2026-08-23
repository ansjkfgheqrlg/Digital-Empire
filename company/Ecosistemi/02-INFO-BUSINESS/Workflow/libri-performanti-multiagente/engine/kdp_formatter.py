"""
KDP Formatter (PIANO-KDP-67, CP6) — assembla un manoscritto in .docx con formattazione
KDP reale (trim size, margini specchio, font, heading, section break per capitolo,
numero pagina) via python-docx + XML diretto dove l'API alto livello non arriva.

Punto critico (bug ripetuto 2 volte nello zip originale analizzato — vedi
CP-20260805-001 §Audit): il conteggio pagine va SEMPRE fatto rileggendo il file appena
salvato, mai fidandosi della sola costruzione in memoria. `build_manuscript_docx`
ritorna un `PageCountResult.within_target` — l'orchestratore (CP9) deve usarlo per
decidere se tornare a scrivere altri capitoli, mai dichiarare fatto se False.
"""
from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from . import config


@dataclass
class Chapter:
    title: str
    paragraphs: list[str]


@dataclass
class BookManuscript:
    title: str
    author: str
    chapters: list[Chapter]
    subtitle: str = ""
    front_matter_notes: list[str] = field(default_factory=list)
    prologue: list[str] | None = None
    epilogue: list[str] | None = None
    # Pagine che non sono il romanzo: copyright davanti, richiesta di recensione e "Also
    # by" in fondo (2026-08-23). Ognuna e' (titolo, paragrafi). Le costruisce `paratesto`.
    pagine_iniziali: list[tuple[str, list[str]]] = field(default_factory=list)
    pagine_finali: list[tuple[str, list[str]]] = field(default_factory=list)


# Le pagine di contorno vanno nel libro ma NON nel conteggio parole che decide se il
# romanzo e' abbastanza lungo (2026-08-23). Senza questa distinzione, aggiungere il
# copyright e la richiesta di recensione avrebbe regalato ~250 parole al controllo di
# lunghezza: un libro con 250 parole di STORIA in meno sarebbe passato lo stesso. Lo stile
# viene salvato dentro il .docx, quindi la distinzione sopravvive alla rilettura del file —
# che e' come qui si contano le cose, mai dalla struttura in memoria.
PARATESTO_STYLE = "Paratesto"


def _nome_stile(paragrafo) -> str:
    """Il nome dello stile, o stringa vuota. Un .docx riletto puo' avere paragrafi senza
    stile associato: chiederglielo direttamente solleverebbe, e far cadere il conteggio
    parole per un dettaglio di formato sarebbe assurdo."""
    try:
        return paragrafo.style.name or ""
    except Exception:  # noqa: BLE001
        return ""


@dataclass
class PageCountResult:
    word_count: int
    estimated_pages: float
    target_min_words: int
    target_max_words: int
    within_target: bool

    def __str__(self) -> str:
        stato = "OK entro il target" if self.within_target else "FUORI TARGET — servono più capitoli"
        return (f"{self.word_count} parole = {self.estimated_pages} pagine @"
                f"{config.WORDS_PER_PAGE_ESTIMATE}wpp (target {self.target_min_words}-"
                f"{self.target_max_words} parole) — {stato}")


def count_words_and_pages(doc: Document) -> PageCountResult:
    """Conta le parole REALI nel documento (mai un numero dichiarato senza ricontrollo).

    Salta le pagine di contorno (copyright, richiesta di recensione, "Also by"): stanno nel
    libro ma non sono il romanzo, e farle contare vorrebbe dire accettare un romanzo piu'
    corto del minimo."""
    word_count = sum(len(p.text.split()) for p in doc.paragraphs
                     if p.text.strip() and _nome_stile(p) != PARATESTO_STYLE)
    estimated_pages = round(word_count / config.WORDS_PER_PAGE_ESTIMATE, 1)
    within_target = config.TARGET_WORD_COUNT_MIN <= word_count <= config.TARGET_WORD_COUNT_MAX
    return PageCountResult(
        word_count=word_count,
        estimated_pages=estimated_pages,
        target_min_words=config.TARGET_WORD_COUNT_MIN,
        target_max_words=config.TARGET_WORD_COUNT_MAX,
        within_target=within_target,
    )


def _set_trim_and_mirror_margins(section) -> None:
    section.page_width = Inches(config.TRIM_SIZE_INCHES[0])
    section.page_height = Inches(config.TRIM_SIZE_INCHES[1])
    section.left_margin = Inches(config.MARGIN_INSIDE_INCHES)
    section.right_margin = Inches(config.MARGIN_OUTSIDE_INCHES)
    section.top_margin = Inches(config.MARGIN_TOP_INCHES)
    section.bottom_margin = Inches(config.MARGIN_BOTTOM_INCHES)
    # Mirror margins reale: non esposto in API alto livello python-docx, XML diretto.
    sect_pr = section._sectPr
    if sect_pr.find(qn("w:mirrorMargins")) is None:
        sect_pr.append(OxmlElement("w:mirrorMargins"))


def _add_page_number_field(footer) -> None:
    """Campo PAGE reale (Word lo calcola all'apertura), non un numero statico finto."""
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def _style_body(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = config.BODY_FONT_NAME
    normal.font.size = Pt(config.BODY_FONT_SIZE_PT)
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing = 1.25
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = Inches(0.2)

    style_names = [s.name for s in document.styles]
    if "Heading 1" in style_names:
        h1 = document.styles["Heading 1"]
        h1.font.name = config.HEADING_FONT_NAME
        h1.font.size = Pt(config.HEADING_FONT_SIZE_PT)
        h1.font.bold = True
        hf = h1.paragraph_format
        hf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hf.space_before = Pt(60)
        hf.space_after = Pt(60)


def _crea_stile_paratesto(document: Document):
    """Stile dedicato alle pagine di contorno, per poterle ESCLUDERE dal conteggio parole
    anche dopo aver riletto il file da disco."""
    from docx.enum.style import WD_STYLE_TYPE

    esistenti = [s.name for s in document.styles]
    if PARATESTO_STYLE in esistenti:
        return document.styles[PARATESTO_STYLE]
    stile = document.styles.add_style(PARATESTO_STYLE, WD_STYLE_TYPE.PARAGRAPH)
    stile.base_style = document.styles["Normal"]
    stile.font.name = config.BODY_FONT_NAME
    stile.font.size = Pt(config.BODY_FONT_SIZE_PT - 1)
    pf = stile.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.first_line_indent = Inches(0)
    pf.space_after = Pt(10)
    return stile


def build_manuscript_docx(manuscript: BookManuscript, output_path: Path) -> PageCountResult:
    """Assembla il manoscritto in .docx con formattazione KDP applicata per davvero.
    Ritorna il conteggio pagine REALE, riletto dal file appena salvato su disco."""
    document = Document()
    _style_body(document)
    _set_trim_and_mirror_margins(document.sections[0])
    _add_page_number_field(document.sections[0].footer)

    _crea_stile_paratesto(document)

    document.add_heading(manuscript.title, level=0)
    if manuscript.subtitle:
        document.add_paragraph(manuscript.subtitle)
    document.add_paragraph(f"by {manuscript.author}")
    for note in manuscript.front_matter_notes:
        document.add_paragraph(note)

    def _new_chapter_section() -> None:
        section = document.add_section(WD_SECTION.NEW_PAGE)
        _set_trim_and_mirror_margins(section)

    def _pagina_paratesto(titolo: str, paragrafi: list[str]) -> None:
        _new_chapter_section()
        if titolo:
            intestazione = document.add_paragraph(titolo, style=PARATESTO_STYLE)
            intestazione.alignment = WD_ALIGN_PARAGRAPH.CENTER
            intestazione.runs[0].bold = True
        for testo in paragrafi:
            p = document.add_paragraph(testo, style=PARATESTO_STYLE)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for titolo, paragrafi in manuscript.pagine_iniziali:
        _pagina_paratesto(titolo, paragrafi)

    if manuscript.prologue:
        _new_chapter_section()
        document.add_heading("Prologue", level=1)
        for para in manuscript.prologue:
            document.add_paragraph(para)

    for chapter in manuscript.chapters:
        _new_chapter_section()
        document.add_heading(chapter.title, level=1)
        for para in chapter.paragraphs:
            document.add_paragraph(para)

    if manuscript.epilogue:
        _new_chapter_section()
        document.add_heading("Epilogue", level=1)
        for para in manuscript.epilogue:
            document.add_paragraph(para)

    for titolo, paragrafi in manuscript.pagine_finali:
        _pagina_paratesto(titolo, paragrafi)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))

    # Non fidarsi della costruzione in memoria: rileggere il file salvato e ricontare.
    reloaded = Document(str(output_path))
    return count_words_and_pages(reloaded)


if __name__ == "__main__":
    import sys
    import tempfile

    print("=== CP6 self-test: 2 casi, uno SOTTO target (deve fallire la validazione) "
          "e uno DENTRO target (deve passare) ===\n")

    def _fake_chapters(n_chapters: int, words_per_chapter: int) -> list[Chapter]:
        filler = " ".join(["parola"] * words_per_chapter)
        return [Chapter(title=f"Chapter {i+1}", paragraphs=[filler, filler]) for i in range(n_chapters)]

    with tempfile.TemporaryDirectory() as tmp:
        tmp_path = Path(tmp)

        # Caso 1: 9 capitoli x ~220 parole = ~2000 parole totali -> SOTTO target (replica
        # esatta del bug reale trovato nello zip: "9 capitoli, ~2000 parole, non 120 pagine")
        short_manuscript = BookManuscript(
            title="Test Corto (deve fallire)", author="Test", chapters=_fake_chapters(9, 220)
        )
        short_result = build_manuscript_docx(short_manuscript, tmp_path / "corto.docx")
        print(f"[Caso 1 - corto] {short_result}")
        assert not short_result.within_target, "ERRORE: doveva essere fuori target e non lo è"
        print("  -> corretto: rilevato fuori target, come atteso\n")

        # Caso 2: 12 capitoli x 2 paragrafi x 1500 parole = ~36000 parole -> DENTRO target (120 pagine)
        long_manuscript = BookManuscript(
            title="Test Lungo (deve passare)", author="Test", chapters=_fake_chapters(12, 1500)
        )
        long_result = build_manuscript_docx(long_manuscript, tmp_path / "lungo.docx")
        print(f"[Caso 2 - lungo] {long_result}")
        assert long_result.within_target, "ERRORE: doveva essere dentro target e non lo è"
        print("  -> corretto: rilevato dentro target, come atteso\n")

        # Verifica XML reale: mirrorMargins e campo PAGE presenti nel file salvato
        reloaded = Document(str(tmp_path / "lungo.docx"))
        sect_pr = reloaded.sections[0]._sectPr
        has_mirror = sect_pr.find(qn("w:mirrorMargins")) is not None
        footer_xml = reloaded.sections[0].footer.paragraphs[0]._p.xml
        has_page_field = "PAGE" in footer_xml
        print(f"[Verifica XML] mirrorMargins presente: {has_mirror}")
        print(f"[Verifica XML] campo PAGE presente nel footer: {has_page_field}")
        assert has_mirror and has_page_field, "ERRORE: formattazione KDP non applicata nell'XML reale"

    print("\nCP6 self-test: TUTTO VERIFICATO OK")
    sys.exit(0)
